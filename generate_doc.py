from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('The Listening Theory in Listen by Wipfler & Schore', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Summary of the listening theory from Listen: Five Simple Tools to Meet Your Everyday Parenting Challenges by Patty Wipfler & Tosha Schore.'
)

doc.add_paragraph('')

# Section 1
doc.add_heading('1. The "Connection Tank" Model', level=1)
doc.add_paragraph(
    "Children's behavior is directly tied to how connected they feel. When the connection tank is full, children cooperate, learn, and play well. "
    "When it's depleted — by stress, transitions, fear, or unmet needs — behavior breaks down. Listening is the primary way to refill it."
)

# Section 2
doc.add_heading('2. Emotions Are Not the Problem — Blocked Emotions Are', level=1)
doc.add_paragraph(
    'Wipfler and Schore argue that the body has a natural recovery system:'
)
bullet_points = [
    'Crying releases grief and stress hormones',
    'Tantrums release frustration',
    'Laughter releases fear',
    'Trembling and sweating are part of healing',
]
for point in bullet_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph(
    'When adults interrupt or suppress these releases (with distractions, bribes, or punishment), the emotion doesn\'t disappear — '
    'it accumulates as a layer of hurt that eventually drives difficult behavior.'
)

# Section 3
doc.add_heading('3. Staylistening — The Core Listening Tool', level=1)
doc.add_paragraph(
    'The practical expression of this theory is Staylistening: the parent stays physically close, maintains warm eye contact, '
    'and listens without trying to fix, stop, or reason away the child\'s feelings. This is not passive — it is an active, attuned presence.'
)
doc.add_paragraph(
    'The key distinction: children crying alone (e.g., "cry it out") do not show the same healing benefits. '
    'The empathic listener is essential. The parent\'s calm, connected presence signals safety to the child\'s nervous system, '
    'allowing the emotional release to complete.'
)

# Section 4
doc.add_heading('4. The Neuroscience Basis', level=1)
doc.add_paragraph('The theory draws on brain science:')
neuro_points = [
    'The limbic system (the brain\'s emotional hub) is wired for connection from birth',
    'Co-regulation — a parent\'s calm nervous system helping to settle a child\'s activated one — is how children learn to eventually self-regulate',
    'When a parent attunes ("I see you\'re upset, I\'m here"), new neural circuitry is built, and recovery takes place',
]
for point in neuro_points:
    doc.add_paragraph(point, style='List Bullet')

# Section 5
doc.add_heading('5. After the Release — The Payoff', level=1)
doc.add_paragraph(
    'Once the emotional storm fully passes (with a listening parent present), children consistently return to a calmer, more cooperative, '
    'more flexible state. The book argues this is not a coincidence — it is the natural result of the nervous system completing its healing cycle.'
)

# Summary
doc.add_heading('In Short', level=1)
p = doc.add_paragraph()
run = p.add_run(
    "The theory is: children's difficult behavior is stored emotional pain, and a listening parent's presence is the mechanism that lets that pain heal. "
    "Listening isn't passive kindness — it's the active ingredient in emotional recovery."
)
run.italic = True

doc.add_paragraph('')

# Resources
doc.add_heading('Sources & Resources', level=1)

resources = [
    ('Hand in Hand Parenting – Staylistening Science',
     'https://www.handinhandparenting.org/2017/01/science-behind-hand-hand-parenting-tool-staylistening/'),
    ('LinkedIn – Science Behind Staylistening',
     'https://www.linkedin.com/pulse/science-behind-hand-parenting-tool-staylistening-lyra-l-estrange-phd'),
    ('Hand in Hand Parenting – Listen Book',
     'https://www.handinhandparenting.org/listen-five-simple-tools'),
    ('Sunshine Parenting Podcast – Ep. 31',
     'https://sunshine-parenting.com/ep-31-listen/'),
]

for label, url in resources:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{label}: ')
    # Add hyperlink
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    # Add hyperlink relationship
    part = doc.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = url
    new_run.append(text_elem)
    hyperlink.append(new_run)
    p._p.append(hyperlink)

output_path = '/home/user/helen-calabria/Listen_Parenting_Summary.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
